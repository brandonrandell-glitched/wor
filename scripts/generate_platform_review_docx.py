#!/usr/bin/env python3
"""Generate a colleague-facing Word document describing the GTM Agent Ecosystem."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
SAMPLES_PATH = ROOT / "output" / "platform-review-samples.json"
OUTPUT_PATH = ROOT / "output" / "GTM-Agent-Ecosystem-Platform-Review.docx"


def _load_samples() -> dict:
    if SAMPLES_PATH.exists():
        return json.loads(SAMPLES_PATH.read_text(encoding="utf-8"))
    return {}


def _add_code(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph(line)
        p.style = "No Spacing"
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)


def _add_json_sample(doc: Document, data: object, max_lines: int = 35) -> None:
    text = json.dumps(data, indent=2, ensure_ascii=False, default=str)
    lines = text.splitlines()
    if len(lines) > max_lines:
        lines = lines[: max_lines - 2] + ["  ...", "}"]
    _add_code(doc, "\n".join(lines))


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val


def build_review_doc(samples: dict) -> Path:
    doc = Document()
    today = date.today().isoformat()

    title = doc.add_heading("GTM Agent Ecosystem — Platform Review", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub = doc.add_paragraph(
        f"Cisco Canada · Partner Security AE · Generated {today}\n"
        "Purpose: colleague review of architecture, MCP tools, workflows, narrative routing, "
        "and sample outputs. All partner/deal data is SAMPLE unless noted."
    )
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "This repository is a multi-workflow GTM assistant platform for Cisco sellers and "
        "partner-facing motions. It combines three customer-facing document workflows "
        "(proposal, discovery prep, competitive brief), seven Cursor MCP servers, a web UI, "
        "and a security-led / platform-capable narrative engine for Tier 2 disti-led selling."
    )
    doc.add_paragraph(
        "Design principle: no live API credentials. All Salesforce, SalesConnect, and partner "
        "data runs from public fixtures and sample JSON until compliance clears real data."
    )

    inv = samples.get("tool_inventory", {})
    total_tools = sum(inv.values()) if inv else 0
    doc.add_paragraph(f"Total MCP tools registered: {total_tools} across 7 servers.")

    doc.add_heading("2. Architecture", level=1)
    _add_code(
        doc,
        """agents/router.py              ← workflow picker (proposal / discovery / competitive)
agents/proposal_assistant.py    ← full proposal intake Q&A
agents/discovery_assistant.py   ← discovery prep intake
agents/competitive_assistant.py ← competitive brief intake
lib/gtm_tools.py                ← shared tools (Salesforce fixtures, pain points, platform story)
lib/public_content.py           ← public Cisco proof points, products, competitive, offers
story_library/                  ← document generators (.docx / .pptx)
web/app.py                      ← web UI with workflow picker
mcp_servers/                    ← GTM MCP (salesforce, proposal-tools, salesconnect)
mcp-framework/                  ← Partner-Ops MCP suite (continuity, cisco-content, partner-ops, sni-scorer)""",
    )

    doc.add_heading("3. Customer-facing workflows", level=1)
    _add_table(
        doc,
        ["Workflow", "Agent", "Output", "How to run"],
        [
            ["Build Proposal", "ProposalAssistant", "Word or PPT proposal", "python -m agents.cli \"Acme\" --workflow proposal --generate"],
            ["Discovery Prep", "DiscoveryAssistant", "Discovery brief (.docx)", "python -m agents.cli \"Acme\" --workflow discovery --generate"],
            ["Competitive Brief", "CompetitiveAssistant", "Competitive brief (.docx)", "python -m agents.cli \"Acme\" --workflow competitive --generate"],
        ],
    )
    doc.add_paragraph(
        "Web UI: python -m web.app → http://127.0.0.1:8080"
    )

    docs = samples.get("generated_documents", {})
    if docs:
        doc.add_paragraph("Sample documents generated for this review (Acme Financial Services demo scenario):")
        for kind, path in docs.items():
            doc.add_paragraph(f"• {kind.title()}: {path}", style="List Bullet")

    doc.add_heading("4. MCP server inventory", level=1)
    tool_rows = [[server, str(count)] for server, count in inv.items()]
    _add_table(doc, ["MCP server", "Tool count"], tool_rows)

    doc.add_heading("4.1 GTM MCP servers (mcp_servers/)", level=2)
    _add_table(
        doc,
        ["Server", "Tools", "Purpose"],
        [
            ["salesforce", "get_customer_context, suggest_opportunities", "Demo account context from fixtures/salesforce_customer.json"],
            ["proposal-tools", "extract_pain_points, recommend_products, recommend_platform_story", "Pain → product mapping with platform routing"],
            ["salesconnect", "get_proof_point, search_competitive, get_offer_detail", "Public Cisco content from fixtures/public_content.json"],
        ],
    )

    doc.add_heading("4.2 Partner-Ops MCP suite (mcp-framework/)", level=2)
    _add_table(
        doc,
        ["Server", "Tools", "Purpose"],
        [
            ["cisco-content", "11 tools incl. build_platform_brief, get_platform_story", "Audience profiles, messaging library, platform narrative drafting"],
            ["partner-ops", "9 tools incl. platform_view, whats_due, qbr_package", "Sample partner book, pipeline, cadence (SAMPLE data)"],
            ["sni-scorer", "get_questionnaire, score_assessment, sprint_status", "Secure Networking Index 30-min audit scoring"],
            ["continuity", "save_recap, resume, log_decision, list_projects", "Session continuity across agent conversations"],
        ],
    )

    doc.add_heading("5. Narrative routing model", level=1)
    doc.add_paragraph(
        "Inbound is expected security-first, but networking, data centre, and collaboration "
        "are first-class entry points. The platform never catalogue-sells across pillars."
    )
    _add_table(
        doc,
        ["Situation", "Story mode", "Behaviour"],
        [
            ["CISO-only sponsor, security budget, Qualified security motion", "security-only", "SME security depth — no pull-through"],
            ["Security inbound; refresh / sprawl / hybrid / AI signals", "security-led", "Security door → platform thread when earned"],
            ["Network / DC / collab RFP or renewal first", "pillar-first", "SME on that pillar; optional pivot when policy/compliance surfaces"],
            ["Security opened door; single-pillar evaluation", "pillar-deep", "Tie-back to security trigger, then SME depth"],
        ],
    )

    doc.add_heading("6. Sample data sources", level=1)
    _add_table(
        doc,
        ["Fixture", "Purpose"],
        [
            ["fixtures/public_content.json", "Products, proof points, competitive, offers"],
            ["fixtures/salesforce_customer.json", "Acme Financial Services demo scenario"],
            ["fixtures/proposal_tools_data.json", "Pain-point patterns, platform pull-through"],
            ["mcp-framework/servers/data/partner_ops.json", "Fictional Tier 1/2 partners and deals (SAMPLE)"],
            ["mcp-framework/servers/data/cisco_content.json", "Audiences, platform model, messaging library"],
        ],
    )

    doc.add_heading("7. Live sample outputs", level=1)
    doc.add_paragraph(
        "The following JSON excerpts were captured by running the built-in logic against "
        "the Acme Financial Services demo scenario on the date of this document."
    )

    sections = [
        ("7.1 Pain point extraction", "extract_pain_points"),
        ("7.2 Product recommendations", "recommend_products"),
        ("7.3 Platform story — security entry", "recommend_platform_story_security"),
        ("7.4 Platform story — network-first entry", "recommend_platform_story_network"),
        ("7.5 Salesforce customer context (SAVM stripped)", "salesforce_context"),
        ("7.6 SNI audit score (illustrative)", "sni_score_sample"),
        ("7.7 Partner pipeline — platform view", "platform_view"),
        ("7.8 Weekly cadence — what's due", "whats_due"),
        ("7.9 Platform story routing (network entry)", "platform_story_network"),
    ]
    for heading, key in sections:
        if key in samples:
            doc.add_heading(heading, level=2)
            _add_json_sample(doc, samples[key])

    doc.add_heading("8. Platform brief logic (build_platform_brief)", level=1)
    doc.add_paragraph(
        "The cisco-content MCP assembles a full drafting package: audience profile, routing rules, "
        "messaging, narrative prompt, and format prompt. Three modes demonstrated below."
    )
    for label, key in [
        ("Security-led (default)", "build_platform_brief_security_led"),
        ("Security-only", "build_platform_brief_security_only"),
        ("Pillar-first (networking inbound)", "build_platform_brief_pillar_first"),
    ]:
        if key in samples:
            doc.add_heading(label, level=2)
            brief = samples[key]
            summary = {
                "story_mode": brief.get("story_mode"),
                "entry_pillar": brief.get("entry_pillar"),
                "instruction": brief.get("instruction"),
                "routing": brief.get("routing"),
                "security_messaging_count": len(brief.get("security_messaging", [])),
                "platform_threads_count": len(brief.get("platform_threads", [])),
                "pull_through_count": len(brief.get("pull_through_plays", [])),
                "pillar_content_count": len(brief.get("pillar_content", [])),
                "pivot_messaging_count": len(brief.get("pivot_messaging", [])),
            }
            _add_json_sample(doc, summary, max_lines=25)

    doc.add_heading("9. Workflow handoffs", level=1)
    doc.add_paragraph(
        "Workflows share lib/gtm_tools.py and can pass confirmed JSON between motions:"
    )
    for item in [
        "Discovery → Proposal: pain points and technologies from discovery JSON",
        "Competitive → Proposal: Competitors field customizes competitive section",
        "Proposal → Competitive: confirmed technologies seed competitive brief",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("10. Test results", level=1)
    pytest_summary = samples.get("pytest_summary", "Not run")
    doc.add_paragraph(f"pytest tests/: {pytest_summary.splitlines()[-1] if pytest_summary else 'N/A'}")
    doc.add_paragraph(
        "Run locally: .venv/bin/python -m pytest tests/ -v"
    )

    doc.add_heading("11. How to demo for a colleague", level=1)
    steps = [
        "Open repo in Cursor — confirm 7 MCP servers green in MCP panel (.cursor/mcp.json).",
        "Run web UI: python -m web.app and walk through workflow picker for Acme Financial Services.",
        "In Cursor chat, invoke MCP tools: get_customer_context, extract_pain_points, recommend_platform_story.",
        "Partner-Ops lens: whats_due(weekly), platform_view, score_assessment for SNI audit.",
        "Content engine: build_platform_brief(audience='tier2-partners', deliverable='one-pager', story_mode='pillar-first', pillar='networking').",
        "Open generated sample .docx files in output/ folder (proposal, discovery, competitive, this review).",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("12. What is not in scope (public lens)", level=1)
    for item in [
        "Real partner PVI, deal values, or internal Cisco links",
        "Live Salesforce, SalesConnect, or OAuth integrations",
        "Official SNI questionnaire text (sample questions only)",
        "Production deployment or CI/CD pipeline",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("")
    footer = doc.add_paragraph("Cisco Canada Partner Growth · GTM Agent Ecosystem · Platform Review")
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    OUTPUT_PATH.parent.mkdir(exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


def main() -> None:
    samples = _load_samples()
    path = build_review_doc(samples)
    print(f"Review document saved to: {path}")


if __name__ == "__main__":
    main()
